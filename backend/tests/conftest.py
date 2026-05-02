from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

FIXTURES_DIR = Path(__file__).parent / "fixtures"
TEST_TEXT = "Hello from test fixture. Python is a programming language."

# Ukrainian IT text: ~150 words, mentions multiple technologies that survive translation.
_UA_IT_TEXT = """\
Розробка сучасних програмних систем базується на ефективному використанні технологій та методологій
управління проектами. Python є однією з найпопулярніших мов програмування для веб-розробки,
автоматизації та науки про дані. Фреймворк FastAPI дозволяє будувати швидкі та надійні API
з автоматичною генерацією документації OpenAPI. Для зберігання реляційних даних використовується
PostgreSQL — потужна об'єктно-реляційна база даних. Redis застосовується для кешування
та реалізації черг повідомлень. Docker спрощує контейнеризацію та розгортання додатків
у різних середовищах. Kubernetes забезпечує оркестрацію контейнерів та масштабування
мікросервісних архітектур. Git є стандартом системи контролю версій у командній розробці.
Хмарні платформи AWS та Azure надають гнучку інфраструктуру для сучасних проектів.
Методологія Agile та практики DevOps допомагають командам швидше постачати якісне програмне
забезпечення завдяки автоматизації тестування та безперервній інтеграції через GitHub Actions.\
"""


@pytest.fixture(scope="session")
def fixtures_dir() -> Path:
    """Create fixture files once per test session and return their directory."""
    FIXTURES_DIR.mkdir(exist_ok=True)

    txt_path = FIXTURES_DIR / "sample.txt"
    txt_path.write_text(TEST_TEXT, encoding="utf-8")

    docx_path = FIXTURES_DIR / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph(TEST_TEXT)
    doc.save(str(docx_path))

    pdf_path = FIXTURES_DIR / "sample.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, TEST_TEXT)
    pdf.output(str(pdf_path))

    # Ukrainian IT text .docx used by test_document_processor
    ua_path = FIXTURES_DIR / "ua_it_sample.docx"
    ua_doc = DocxDocument()
    ua_doc.add_heading("Розробка програмного забезпечення", level=1)
    for paragraph in _UA_IT_TEXT.split("\n"):
        if paragraph.strip():
            ua_doc.add_paragraph(paragraph.strip())
    ua_doc.save(str(ua_path))

    return FIXTURES_DIR


@pytest.fixture(scope="session")
def ua_it_docx_path(fixtures_dir: Path) -> Path:
    """Return path to the Ukrainian IT text .docx fixture."""
    return fixtures_dir / "ua_it_sample.docx"
